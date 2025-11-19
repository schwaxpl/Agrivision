"""
Loader pour charger et traiter des documents Word (.docx) avec LangChain.
"""

import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from langchain_core.documents import Document
import docx
import logging


class WordLoader:
    """
    Loader pour charger des documents Word (.docx) et les convertir en documents LangChain.
    """
    
    def __init__(self, 
                 extract_tables: bool = True,
                 preserve_formatting: bool = False):
        """
        Initialise le loader Word.
        
        Args:
            extract_tables: Si True, extrait le contenu des tableaux
            preserve_formatting: Si True, tente de préserver la mise en forme basique
        """
        self.extract_tables = extract_tables
        self.preserve_formatting = preserve_formatting
        self.logger = logging.getLogger(__name__)
    
    def load(self, file_path: str) -> List[Document]:
        """
        Charge un fichier Word et retourne une liste de documents LangChain.
        
        Args:
            file_path: Chemin vers le fichier .docx
            
        Returns:
            Liste contenant un document LangChain avec le contenu extrait
            
        Raises:
            FileNotFoundError: Si le fichier n'existe pas
            Exception: Si l'extraction échoue
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Fichier non trouvé: {file_path}")
        
        if not file_path.lower().endswith('.docx'):
            raise ValueError(f"Le fichier doit avoir l'extension .docx: {file_path}")
        
        try:
            # Chargement du document Word
            doc = docx.Document(file_path)
            
            # Extraction du contenu
            content_parts = []
            
            # Extraction des paragraphes
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    if self.preserve_formatting:
                        # Préservation basique du formatage
                        if paragraph.style.name.startswith('Heading'):
                            level = self._get_heading_level(paragraph.style.name)
                            content_parts.append(f"{'#' * level} {text}")
                        else:
                            content_parts.append(text)
                    else:
                        content_parts.append(text)
            
            # Extraction des tableaux si demandé
            if self.extract_tables:
                for table in doc.tables:
                    table_content = self._extract_table_content(table)
                    if table_content:
                        content_parts.append(table_content)
            
            # Assemblage du contenu final
            full_content = "\n\n".join(content_parts)
            
            # Métadonnées du document
            metadata = self._extract_metadata(file_path, doc)
            
            # Création du document LangChain
            document = Document(
                page_content=full_content,
                metadata=metadata
            )
            
            self.logger.info(f"Document Word chargé: {file_path} ({len(full_content)} caractères)")
            
            return [document]
            
        except Exception as e:
            self.logger.error(f"Erreur lors du chargement du fichier Word {file_path}: {e}")
            raise Exception(f"Impossible de charger le fichier Word: {str(e)}")
    
    def load_directory(self, directory_path: str, 
                      pattern: str = "*.docx",
                      recursive: bool = True) -> List[Document]:
        """
        Charge tous les fichiers Word d'un répertoire.
        
        Args:
            directory_path: Chemin vers le répertoire
            pattern: Pattern pour filtrer les fichiers (défaut: "*.docx")
            recursive: Si True, recherche récursivement dans les sous-répertoires
            
        Returns:
            Liste de documents LangChain
        """
        directory = Path(directory_path)
        
        if not directory.exists():
            raise FileNotFoundError(f"Répertoire non trouvé: {directory_path}")
        
        # Recherche des fichiers
        if recursive:
            files = list(directory.rglob(pattern))
        else:
            files = list(directory.glob(pattern))
        
        documents = []
        failed_files = []
        
        for file_path in files:
            try:
                docs = self.load(str(file_path))
                documents.extend(docs)
                print(f"✅ Fichier chargé: {file_path.name}")
            except Exception as e:
                failed_files.append((str(file_path), str(e)))
                print(f"❌ Erreur avec {file_path.name}: {e}")
        
        if failed_files:
            print(f"\n⚠️  {len(failed_files)} fichiers ont échoué:")
            for file_path, error in failed_files:
                print(f"  - {Path(file_path).name}: {error}")
        
        print(f"\n📊 Total: {len(documents)} documents chargés depuis {len(files)} fichiers")
        return documents
    
    def _extract_table_content(self, table) -> str:
        """
        Extrait le contenu d'un tableau Word.
        
        Args:
            table: Tableau docx
            
        Returns:
            Contenu du tableau formaté en texte
        """
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            
            if any(cell.strip() for cell in row_data):  # Ignorer les lignes vides
                table_data.append(" | ".join(row_data))
        
        if table_data:
            return "\n[TABLE]\n" + "\n".join(table_data) + "\n[/TABLE]\n"
        return ""
    
    def _get_heading_level(self, style_name: str) -> int:
        """
        Détermine le niveau de titre à partir du nom du style.
        
        Args:
            style_name: Nom du style Word
            
        Returns:
            Niveau de titre (1-6)
        """
        if 'Heading' in style_name:
            try:
                # Extrait le numéro du style (ex: "Heading 1" -> 1)
                level = int(style_name.split()[-1])
                return min(level, 6)  # Limite à 6 niveaux
            except (ValueError, IndexError):
                return 1
        return 1
    
    def _extract_metadata(self, file_path: str, doc) -> Dict[str, Any]:
        """
        Extrait les métadonnées du document Word.
        
        Args:
            file_path: Chemin vers le fichier
            doc: Document docx
            
        Returns:
            Dictionnaire des métadonnées
        """
        file_path_obj = Path(file_path)
        
        metadata = {
            "source": str(file_path_obj),
            "file_name": file_path_obj.name,
            "file_type": "docx",
            "file_size": file_path_obj.stat().st_size if file_path_obj.exists() else 0,
        }
        
        # Propriétés du document Word
        try:
            core_props = doc.core_properties
            if hasattr(core_props, 'title') and core_props.title:
                metadata["title"] = core_props.title
            if hasattr(core_props, 'author') and core_props.author:
                metadata["author"] = core_props.author
            if hasattr(core_props, 'subject') and core_props.subject:
                metadata["subject"] = core_props.subject
            if hasattr(core_props, 'created') and core_props.created:
                metadata["created_date"] = core_props.created.isoformat()
            if hasattr(core_props, 'modified') and core_props.modified:
                metadata["modified_date"] = core_props.modified.isoformat()
        except Exception as e:
            self.logger.warning(f"Impossible d'extraire les propriétés du document: {e}")
        
        # Statistiques du contenu
        try:
            metadata["paragraphs_count"] = len(doc.paragraphs)
            metadata["tables_count"] = len(doc.tables)
        except Exception as e:
            self.logger.warning(f"Impossible d'extraire les statistiques: {e}")
        
        return metadata
    
    def get_document_info(self, file_path: str) -> Dict[str, Any]:
        """
        Retourne des informations sur le document sans le charger complètement.
        
        Args:
            file_path: Chemin vers le fichier .docx
            
        Returns:
            Dictionnaire avec les informations du document
        """
        try:
            doc = docx.Document(file_path)
            
            info = {
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "estimated_pages": max(1, len(doc.paragraphs) // 10),  # Estimation grossière
            }
            
            # Propriétés du document
            try:
                core_props = doc.core_properties
                if core_props.title:
                    info["title"] = core_props.title
                if core_props.author:
                    info["author"] = core_props.author
                if core_props.subject:
                    info["subject"] = core_props.subject
            except Exception:
                pass
            
            return info
            
        except Exception as e:
            return {"error": str(e), "file_path": file_path}


def create_word_loader(extract_tables: bool = True, 
                      preserve_formatting: bool = False) -> WordLoader:
    """
    Factory function pour créer un WordLoader avec des paramètres spécifiques.
    
    Args:
        extract_tables: Si True, extrait le contenu des tableaux
        preserve_formatting: Si True, préserve la mise en forme basique
        
    Returns:
        Instance de WordLoader configurée
    """
    return WordLoader(
        extract_tables=extract_tables,
        preserve_formatting=preserve_formatting
    )